#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
更新成本管理-成本总览页面需求文档 v2.2
补充导入弹窗和手动添加弹窗的需求内容
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_elm = parse_xml(r'<w:%s w:val="%s" w:sz="%s" w:space="0" w:color="%s" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>' % (
                edge, kwargs[edge]["val"], kwargs[edge]["sz"], kwargs[edge]["color"]))
            tcPr.append(edge_elm)

def add_heading_with_style(doc, text, level=1, color=None):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(color[0], color[1], color[2])
    return heading

def add_paragraph_with_style(doc, text, bold=False, color=None, size=None):
    """添加带样式的段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor(color[0], color[1], color[2])
    if size:
        run.font.size = Pt(size)
    return p

def add_new_section_marker(doc):
    """添加新增内容标记"""
    p = doc.add_paragraph()
    run = p.add_run("【新增内容】")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
    run.font.size = Pt(11)
    return p

def create_requirements_doc():
    """创建完整的需求文档"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)
    
    # 标题
    title = doc.add_heading('成本管理-成本总览页面需求文档', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(18)
        run.bold = True
    
    # 版本信息
    version_para = doc.add_paragraph()
    version_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = version_para.add_run('版本：v2.2\n日期：2026-02-10\n作者：产品团队')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    
    # 版本历史
    add_heading_with_style(doc, '版本历史', level=1)
    
    # 版本历史表格
    versions = [
        ['v2.2', '2026-02-10', '产品团队', '补充导入弹窗、手动添加弹窗需求（新增）'],
        ['v2.1', '2026-02-10', '产品团队', '修正地区字段为非必填，更新资源详情和计费资源tab页需求'],
        ['v2.0', '2026-02-10', '产品团队', '补充当前月编辑页面和历史月详情页面需求'],
        ['v1.0', '2026-02-07', '产品团队', '初始版本']
    ]
    
    table = doc.add_table(rows=len(versions)+1, cols=4)
    table.style = 'Table Grid'
    
    # 表头
    headers = ['版本', '日期', '作者', '变更内容']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 版本数据
    for i, version_data in enumerate(versions):
        for j, cell_text in enumerate(version_data):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    doc.add_paragraph()
    
    # 1. 页面概述
    add_heading_with_style(doc, '一、页面概述', level=1)
    
    add_heading_with_style(doc, '1.1 页面名称', level=2)
    add_paragraph_with_style(doc, '成本管理-成本总览')
    
    add_heading_with_style(doc, '1.2 页面路径', level=2)
    add_paragraph_with_style(doc, '/cost-overview.html')
    
    add_heading_with_style(doc, '1.3 页面目的', level=2)
    add_paragraph_with_style(doc, '提供成本数据的集中展示、查询、管理和维护功能，支持多维度筛选、批量操作、数据导入导出等能力。')
    
    add_heading_with_style(doc, '1.4 目标用户', level=2)
    add_paragraph_with_style(doc, '• 成本管理员\n• 财务人员\n• 运维管理人员')
    
    doc.add_page_break()
    
    # 2. 功能模块
    add_heading_with_style(doc, '二、功能模块', level=1)
    
    # 2.1 查询筛选区
    add_heading_with_style(doc, '2.1 查询筛选区', level=2)
    
    add_heading_with_style(doc, '2.1.1 账期选择', level=3)
    add_paragraph_with_style(doc, '• 类型：下拉选择\n• 默认值：当前月份\n• 选项：2025-12、2026-01、2026-02\n• 交互：切换月份后自动刷新列表数据', bold=False)
    
    add_heading_with_style(doc, '2.1.2 搜索框', level=3)
    add_paragraph_with_style(doc, '• 占位符：搜索成本项名称/编码\n• 支持实时搜索或回车搜索\n• 支持模糊匹配', bold=False)
    
    add_heading_with_style(doc, '2.1.3 高级筛选', level=3)
    add_paragraph_with_style(doc, '• 成本分类：多选下拉（设备费用、软件平台、网络带宽、运营费用、基础设施）\n• 成本池：多选下拉（通算服务器、智算服务器、存储、安全、网络、机房环境）\n• 公摊池：多选下拉（基础资源、容器、数据库、网络）\n• 公摊类型：多选下拉（公摊成本、非公摊成本）\n• 操作：点击"筛选"按钮应用筛选条件，点击"重置"清空所有筛选条件', bold=False)
    
    # 2.2 操作按钮区
    add_heading_with_style(doc, '2.2 操作按钮区', level=2)
    
    add_heading_with_style(doc, '2.2.1 导入按钮', level=3)
    add_paragraph_with_style(doc, '• 按钮样式：默认样式（灰色边框）\n• 按钮文字：导入\n• 交互：点击打开批量导入弹窗\n• 历史月限制：历史月份（非当前月）时按钮禁用', bold=False)
    
    add_heading_with_style(doc, '2.2.2 手动添加按钮', level=3)
    add_paragraph_with_style(doc, '• 按钮样式：默认样式（灰色边框）\n• 按钮文字：手动添加\n• 交互：点击打开手动添加弹窗\n• 历史月限制：历史月份（非当前月）时按钮禁用', bold=False)
    
    add_heading_with_style(doc, '2.2.3 导出按钮', level=3)
    add_paragraph_with_style(doc, '• 按钮样式：默认样式（灰色边框）\n• 按钮文字：导出\n• 交互：点击导出当前筛选条件下的所有数据为Excel文件\n• 导出字段：与列表显示字段一致', bold=False)
    
    add_heading_with_style(doc, '2.2.4 批量删除按钮', level=3)
    add_paragraph_with_style(doc, '• 按钮样式：危险样式（红色边框）\n• 按钮文字：批量删除\n• 交互：\n  - 未选择数据时：提示"请先选择要删除的数据"\n  - 已选择数据时：弹出确认对话框，确认后删除选中数据\n• 历史月限制：历史月份（非当前月）时按钮禁用', bold=False)
    
    doc.add_page_break()
    
    # 2.3 数据列表区
    add_heading_with_style(doc, '2.3 数据列表区', level=2)
    
    add_heading_with_style(doc, '2.3.1 列表字段', level=3)
    
    # 字段数据
    fields = [
        ['1', '选择框', '复选框', '否', '支持跨页多选'],
        ['2', '成本编码', '文本', '是', '如：DEV-001、OPS-019'],
        ['3', '成本分类', '文本', '是', '设备费用/软件平台/网络带宽/运营费用/基础设施'],
        ['4', '成本池', '文本', '是', '通算服务器/智算服务器/存储/安全/网络'],
        ['5', '成本项', '文本', '是', '点击可进入详情页'],
        ['6', '地域', '文本', '是', '北京一期/北京二期/西安/上海'],
        ['7', '公摊池', '文本', '是', '基础资源/容器/数据库/网络'],
        ['8', '公摊类型', '文本', '是', '公摊成本/非公摊成本'],
        ['9', '数据来源', '文本', '是', '资产系统-B类资产/手动维护'],
        ['10', '描述', '文本', '否', '成本项描述信息'],
        ['11', '单价（元）', '数值', '是', '原始采购单价'],
        ['12', '折旧月限', '数值', '是', '折旧期限（月）'],
        ['13', '数量', '数值', '是', '资源数量'],
        ['14', '单位', '文本', '是', '台/套/Gbps/TB/月'],
        ['15', '账期', '文本', '是', 'YYYY-MM格式'],
        ['16', '月成本', '数值', '是', '本月分摊成本'],
        ['17', '公摊比例', '数值', '是', '公摊比例（%）'],
        ['18', '公摊后月总成本', '数值', '是', '最终成本金额'],
        ['19', '操作', '按钮组', '否', '编辑/删除/详情按钮']
    ]
    
    # 列表字段表格
    table = doc.add_table(rows=len(fields)+1, cols=5)
    table.style = 'Table Grid'
    
    # 表头
    headers = ['序号', '字段名称', '字段类型', '是否可排序', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    for i, field_data in enumerate(fields):
        for j, cell_text in enumerate(field_data):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    add_heading_with_style(doc, '2.3.2 分页功能', level=3)
    add_paragraph_with_style(doc, '• 默认每页显示：10条\n• 可选分页：10/20/50/100条\n• 分页组件：首页/上一页/页码/下一页/末页\n• 总条数显示：共 XX 条', bold=False)
    
    add_heading_with_style(doc, '2.3.3 排序功能', level=3)
    add_paragraph_with_style(doc, '• 支持点击表头进行单列排序\n• 排序方式：升序/降序切换\n• 默认排序：按成本编码升序', bold=False)
    
    add_heading_with_style(doc, '2.3.4 行选择功能', level=3)
    add_paragraph_with_style(doc, '• 支持单行选择（复选框）\n• 支持全选（表头复选框）\n• 支持跨页选择（切换分页保留选中状态）\n• 已选择数量显示在批量删除按钮旁', bold=False)
    
    add_heading_with_style(doc, '2.3.5 操作列按钮', level=3)
    add_paragraph_with_style(doc, '当前月（2026-02）操作列显示：', bold=False)
    add_paragraph_with_style(doc, '• 编辑按钮：点击进入编辑页面\n• 删除按钮：点击删除该行数据（需确认）', bold=False)
    add_paragraph_with_style(doc, '历史月（2025-12、2026-01）操作列显示：', bold=False)
    add_paragraph_with_style(doc, '• 详情按钮：点击进入详情页（只读模式）', bold=False)
    
    doc.add_page_break()
    
    # ==================== 新增内容开始 ====================
    add_new_section_marker(doc)
    
    # 2.4 导入弹窗
    add_heading_with_style(doc, '2.4 批量导入弹窗【新增】', level=2)
    
    add_heading_with_style(doc, '2.4.1 弹窗触发方式', level=3)
    add_paragraph_with_style(doc, '• 点击页面顶部"导入"按钮打开\n• 弹窗标题：批量导入\n• 弹窗尺寸：默认居中，宽度500px\n• 遮罩：显示灰色半透明遮罩，点击遮罩不关闭弹窗', bold=False)
    
    add_heading_with_style(doc, '2.4.2 弹窗字段', level=3)
    
    import_fields = [
        ['1', '上传文件', '文件上传', '是', '支持.xlsx, .xls格式，支持点击上传和拖拽上传'],
        ['2', '数据来源', '下拉选择', '是', '选项：资产系统-B类资产、手动维护（默认选中手动维护）'],
        ['3', '导入选项', '单选', '否', '替换现有数据/追加数据（默认选中追加数据）'],
        ['4', '下载模板', '按钮', '否', '根据所选数据来源下载对应模板']
    ]
    
    # 导入弹窗字段表格
    table = doc.add_table(rows=len(import_fields)+1, cols=5)
    table.style = 'Table Grid'
    
    headers = ['序号', '字段名称', '字段类型', '是否必填', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    for i, field_data in enumerate(import_fields):
        for j, cell_text in enumerate(field_data):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    add_heading_with_style(doc, '2.4.3 文件上传区域', level=3)
    add_paragraph_with_style(doc, '• 样式：虚线边框拖拽区域\n• 提示文字：\n  - 主文案："点击或拖拽文件到此处上传"\n  - 副文案："支持 .xlsx, .xls 文件"\n• 上传按钮："选择文件"按钮，点击唤起文件选择器\n• 拖拽交互：\n  - 拖拽进入：边框变蓝色，背景变浅蓝色\n  - 拖拽离开：恢复原样式\n  - 拖拽释放：读取文件并显示文件名\n• 文件选择后：在区域下方显示"已选择文件: xxx.xlsx"', bold=False)
    
    add_heading_with_style(doc, '2.4.4 数据来源切换逻辑', level=3)
    add_paragraph_with_style(doc, '• 切换数据来源时，模板描述文字同步更新\n• 默认显示："当前模板：手动维护"\n• 切换为资产系统-B类资产后显示："当前模板：资产系统-B类资产"', bold=False)
    
    add_heading_with_style(doc, '2.4.5 下载模板功能', level=3)
    add_paragraph_with_style(doc, '• 按钮文字：下载导入模板\n• 模板类型：\n  - 手动维护模板：成本分类、成本池、成本项、地域、公摊池、公摊类型、单价、折旧月限、数量、单位、描述\n  - 资产系统-B类资产模板：成本分类、成本池、成本项、公摊池、公摊类型、描述\n• 文件格式：CSV格式\n• 文件名：成本数据导入模板-手动维护.csv / 成本数据导入模板-资产系统-B类资产.csv', bold=False)
    
    add_heading_with_style(doc, '2.4.6 弹窗按钮', level=3)
    add_paragraph_with_style(doc, '• 取消按钮：关闭弹窗，不执行任何操作\n• 确认导入按钮：\n  - 校验：未选择文件时提示"请先选择文件"\n  - 成功提示："导入成功"\n  - 关闭弹窗并刷新列表数据', bold=False)
    
    doc.add_page_break()
    
    # 2.5 手动添加弹窗
    add_heading_with_style(doc, '2.5 手动添加弹窗【新增】', level=2)
    
    add_heading_with_style(doc, '2.5.1 弹窗触发方式', level=3)
    add_paragraph_with_style(doc, '• 点击页面顶部"手动添加"按钮打开\n• 弹窗标题：手动添加成本\n• 弹窗尺寸：默认居中，宽度600px\n• 遮罩：显示灰色半透明遮罩，点击遮罩不关闭弹窗', bold=False)
    
    add_heading_with_style(doc, '2.5.2 表单结构', level=3)
    add_paragraph_with_style(doc, '弹窗内包含两个并排的表单容器，根据数据来源切换显示：', bold=False)
    add_paragraph_with_style(doc, '• 资产系统-B类资产表单（id="asset-system-form"）\n• 手动维护表单（id="manual-form"）\n• 默认显示：手动维护表单', bold=False)
    
    add_heading_with_style(doc, '2.5.3 公共字段（两个表单共有）', level=3)
    
    common_fields = [
        ['1', '成本分类', '下拉选择', '是', '选项：设备费用、软件平台、网络带宽、运营费用、基础设施'],
        ['2', '成本池', '下拉选择', '是', '选项：通算服务器、智算服务器、存储、安全、网络、机房环境'],
        ['3', '成本项', '文本输入', '是', '请输入成本项名称']
    ]
    
    # 公共字段表格
    table = doc.add_table(rows=len(common_fields)+1, cols=5)
    table.style = 'Table Grid'
    
    headers = ['序号', '字段名称', '字段类型', '是否必填', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    for i, field_data in enumerate(common_fields):
        for j, cell_text in enumerate(field_data):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    add_heading_with_style(doc, '2.5.4 资产系统-B类资产表单字段', level=3)
    
    asset_fields = [
        ['1', '公摊池', '下拉选择', '是', '选项：基础资源、容器、数据库、网络'],
        ['2', '公摊类型', '下拉选择', '是', '选项：公摊成本、非公摊成本'],
        ['3', '数据来源', '下拉选择', '是', '选项：资产系统-B类资产、手动维护（默认选中资产系统-B类资产）'],
        ['4', '资产分类', '下拉选择', '是', '选项：服务器、存储设备、网络设备、安全设备（当数据来源为资产系统时显示）'],
        ['5', '账期', '月份选择', '是', '系统字段，默认当前月份，不可编辑'],
        ['6', '描述', '文本域', '否', '请输入描述信息，多行文本']
    ]
    
    # 资产系统表单字段表格
    table = doc.add_table(rows=len(asset_fields)+1, cols=5)
    table.style = 'Table Grid'
    
    headers = ['序号', '字段名称', '字段类型', '是否必填', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    for i, field_data in enumerate(asset_fields):
        for j, cell_text in enumerate(field_data):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    add_heading_with_style(doc, '2.5.5 手动维护表单字段', level=3)
    
    manual_fields = [
        ['1', '公摊池', '下拉选择', '是', '选项：基础资源、容器、数据库、网络'],
        ['2', '公摊类型', '下拉选择', '是', '选项：公摊成本、非公摊成本'],
        ['3', '数据来源', '下拉选择', '是', '选项：资产系统-B类资产、手动维护（默认选中手动维护）'],
        ['4', '资产分类', '下拉选择', '是', '选项：服务器、存储设备、网络设备、安全设备（当数据来源为资产系统时显示）'],
        ['5', '地域', '下拉选择', '否', '选项：北京一期、北京二期、西安、上海'],
        ['6', '单价（元）', '数值输入', '是', '请输入单价，支持小数'],
        ['7', '折旧月限', '数值输入', '是', '请输入折旧月限，整数'],
        ['8', '数量', '数值输入', '是', '请输入数量，整数'],
        ['9', '单位', '下拉选择', '是', '选项：台、套、Gbps、TB、月'],
        ['10', '账期', '月份选择', '是', '系统字段，默认当前月份，不可编辑'],
        ['11', '描述', '文本域', '否', '请输入描述信息，多行文本']
    ]
    
    # 手动维护表单字段表格
    table = doc.add_table(rows=len(manual_fields)+1, cols=5)
    table.style = 'Table Grid'
    
    headers = ['序号', '字段名称', '字段类型', '是否必填', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    for i, field_data in enumerate(manual_fields):
        for j, cell_text in enumerate(field_data):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    add_heading_with_style(doc, '2.5.6 数据来源切换逻辑', level=3)
    add_paragraph_with_style(doc, '• 当数据来源选择"资产系统-B类资产"时：\n  - 显示资产分类字段\n  - 隐藏地域、单价、折旧月限、数量、单位字段\n  - 切换显示资产系统-B类资产表单容器', bold=False)
    add_paragraph_with_style(doc, '• 当数据来源选择"手动维护"时：\n  - 隐藏资产分类字段\n  - 显示地域、单价、折旧月限、数量、单位字段\n  - 切换显示手动维护表单容器', bold=False)
    
    add_heading_with_style(doc, '2.5.7 弹窗按钮', level=3)
    add_paragraph_with_style(doc, '• 取消按钮：关闭弹窗，不执行任何操作\n• 保存按钮：\n  - 表单校验：必填字段未填写时提示"请填写必填项"\n  - 成功提示："保存成功"\n  - 关闭弹窗并刷新列表数据\n  - 新增数据默认显示在列表第一行', bold=False)
    
    add_heading_with_style(doc, '2.5.8 表单校验规则', level=3)
    add_paragraph_with_style(doc, '• 必填字段：不能为空\n• 数值字段：必须为正数\n• 单价：支持2位小数\n• 折旧月限/数量：必须为整数\n• 账期：系统自动填充，用户不可修改', bold=False)
    
    doc.add_page_break()
    
    # 3. 交互规则
    add_heading_with_style(doc, '三、交互规则', level=1)
    
    add_heading_with_style(doc, '3.1 历史月数据限制', level=2)
    add_paragraph_with_style(doc, '当账期选择历史月份（非当前月）时：', bold=False)
    add_paragraph_with_style(doc, '• 导入按钮：禁用（灰色，点击无响应）\n• 手动添加按钮：禁用（灰色，点击无响应）\n• 批量删除按钮：禁用（灰色，点击无响应）\n• 列表操作列：只显示"详情"按钮，隐藏"编辑"和"删除"按钮\n• 列表选择框：禁用（不能勾选）', bold=False)
    
    add_heading_with_style(doc, '3.2 列表行点击', level=2)
    add_paragraph_with_style(doc, '• 点击成本项名称：进入详情页\n• 当前月：进入编辑模式（可编辑）\n• 历史月：进入查看模式（只读）', bold=False)
    
    add_heading_with_style(doc, '3.3 成本项编码规则', level=2)
    add_paragraph_with_style(doc, '• 设备费用：DEV-XXX\n• 运营费用：OPS-XXX\n• 软件平台：SFT-XXX\n• 网络带宽：NET-XXX\n• 基础设施：INF-XXX', bold=False)
    
    doc.add_page_break()
    
    # 4. 性能要求
    add_heading_with_style(doc, '四、性能要求', level=1)
    add_paragraph_with_style(doc, '• 列表加载时间：≤ 2秒\n• 搜索响应时间：≤ 1秒\n• 分页切换时间：≤ 1秒\n• 导入数据处理：≤ 5秒（1000条数据）\n• 导出数据处理：≤ 3秒（1000条数据）\n• 弹窗打开时间：≤ 500ms', bold=False)
    
    # 5. 异常处理
    add_heading_with_style(doc, '五、异常处理', level=1)
    add_paragraph_with_style(doc, '• 网络异常：提示"网络连接失败，请检查网络后重试"\n• 服务器异常：提示"服务器繁忙，请稍后重试"\n• 数据加载失败：提示"数据加载失败，请点击刷新重试"\n• 导入文件格式错误：提示"文件格式不正确，请上传.xlsx或.xls格式文件"\n• 导入数据校验失败：提示"第X行数据格式错误：XXX"\n• 必填字段为空：提示"请填写必填项"并聚焦到第一个空字段', bold=False)
    
    # 6. 安全要求
    add_heading_with_style(doc, '六、安全要求', level=1)
    add_paragraph_with_style(doc, '• 敏感操作需二次确认（删除、批量删除）\n• 导入文件大小限制：≤ 10MB\n• 导入文件类型限制：仅允许.xlsx, .xls格式\n• 前端表单校验：防止XSS攻击\n• 数据权限：用户只能操作有权限的成本数据', bold=False)
    
    # 保存文档
    doc.save('e:\\IT运维\\FinOps\\AI coding FinOps\\FinOpsV2.0\\成本管理-成本总览页面需求文档_v2.2.docx')
    print("文档已保存：成本管理-成本总览页面需求文档_v2.2.docx")

if __name__ == '__main__':
    create_requirements_doc()
