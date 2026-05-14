from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    return table

def update_document():
    doc = Document('成本管理-成本总览页面需求文档.docx')
    
    # 添加分页符
    doc.add_page_break()
    
    add_heading(doc, '附录A：列表字段变更记录', level=1)
    
    add_heading(doc, 'A.1 变更说明', level=2)
    add_paragraph(doc, '根据实际业务需求调整，成本总览页面列表字段已进行以下变更：')
    add_paragraph(doc, '• 删除字段：CPU资源总量、MEM资源总量、Disk资源总量')
    add_paragraph(doc, '• 变更原因：资源总量信息移至详情页面展示，列表页仅展示核心成本信息')
    add_paragraph(doc, '• 变更日期：2026-02-10')
    
    add_heading(doc, 'A.2 当前列表字段（变更后）', level=2)
    add_paragraph(doc, '成本总览页面数据列表当前包含以下字段：')
    
    headers = ['序号', '字段名称', '说明', '排序', '宽度']
    rows = [
        ['1', '勾选框', '支持单行勾选和全选', '-', '40px'],
        ['2', '成本编码', '系统生成的唯一编码', '支持', '自适应'],
        ['3', '成本分类', '成本项所属分类', '支持', '自适应'],
        ['4', '成本池', '成本项所属池', '支持', '自适应'],
        ['5', '成本项', '成本项名称（可点击跳转详情）', '支持', '自适应'],
        ['6', '地域', '所属地域', '支持', '自适应'],
        ['7', '公摊池', '公摊池名称', '支持', '自适应'],
        ['8', '公摊类型', '公摊/非公摊', '支持', '自适应'],
        ['9', '描述', '成本项描述', '支持', '自适应'],
        ['10', '单价', '资产单价', '支持', '自适应'],
        ['11', '折旧月限', '折旧期限（月）', '支持', '自适应'],
        ['12', '折旧月价', '每月折旧金额', '支持', '自适应'],
        ['13', '数量', '资产数量', '支持', '自适应'],
        ['14', '单位', '计量单位', '支持', '自适应'],
        ['15', '账期', '所属账期', '支持', '自适应'],
        ['16', '月成本', '当月成本金额', '支持', '自适应'],
        ['17', '公摊比例', '公摊百分比', '支持', '自适应'],
        ['18', '公摊后月总成本', '公摊计算后的成本', '支持', '自适应'],
        ['19', '操作', '编辑/删除/详情按钮', '-', '120px（冻结列）'],
    ]
    add_table(doc, headers, rows)
    
    add_heading(doc, 'A.3 历史列表字段（变更前）', level=2)
    add_paragraph(doc, '变更前的列表字段包含以下22个字段：')
    add_paragraph(doc, '成本编码、成本分类、成本池、成本项、地域、公摊池、公摊类型、描述、单价（元）、折旧月限、折旧月价、数量、单位、账期、月成本（元）、公摊比例、公摊后月总成本、CPU资源总量、MEM资源总量、Disk资源总量、操作')
    add_paragraph(doc, '注：CPU资源总量、MEM资源总量、Disk资源总量三个字段已从列表页移除')
    
    add_heading(doc, 'A.4 导出字段同步变更', level=2)
    add_paragraph(doc, 'Excel/CSV导出功能字段同步调整：')
    add_paragraph(doc, '• 导出字段与列表展示字段保持一致')
    add_paragraph(doc, '• 不包含已删除的CPU、MEM、Disk资源总量字段')
    add_paragraph(doc, '• 导出表头顺序与列表字段顺序一致')
    
    # 保存文档
    output_path = '成本管理-成本总览页面需求文档_v2.2.docx'
    doc.save(output_path)
    print(f"文档更新完成！")
    print(f"保存路径: {output_path}")
    print("\n更新内容：")
    print("1. 添加了附录A：列表字段变更记录")
    print("2. 记录了CPU、MEM、Disk资源总量字段的删除变更")
    print("3. 列出了变更后的19个列表字段")
    print("4. 说明了导出字段同步变更")

if __name__ == '__main__':
    update_document()
