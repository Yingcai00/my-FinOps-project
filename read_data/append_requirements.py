from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, size=10.5):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p

def add_table(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # 数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    return table

def main():
    # 打开现有文档
    doc = Document('成本管理-成本总览页面需求文档.docx')
    
    # 添加分页符
    doc.add_page_break()
    
    # ========== 当前月-编辑页面 ==========
    add_heading(doc, '10. 当前月-编辑页面需求', level=1)
    
    add_heading(doc, '10.1 页面概述', level=2)
    add_paragraph(doc, '当前月编辑页面用于编辑和修改当前账期的成本项信息，支持基本信息维护、资源详情管理和计费资源配置。')
    
    add_heading(doc, '10.2 页面入口', level=2)
    add_paragraph(doc, '• 入口路径：成本总览页面 → 当前月数据行 → 点击"编辑"按钮')
    add_paragraph(doc, '• URL参数：cost-detail.html?id={id}&mode=edit')
    
    add_heading(doc, '10.3 页面布局', level=2)
    add_paragraph(doc, '页面采用左右分栏布局：')
    add_paragraph(doc, '• 左侧：成本项基本信息表单（固定宽度400px）')
    add_paragraph(doc, '• 右侧：Tab切换区域（资源详情/计费资源）')
    
    add_heading(doc, '10.4 基本信息区域', level=2)
    add_paragraph(doc, '基本信息表单字段及规则：', bold=True)
    
    headers = ['字段名称', '控件类型', '是否必填', '编辑状态', '校验规则']
    rows = [
        ['成本编码', '文本输入框', '是', '禁用（系统自动生成）', '只读'],
        ['账期', '月份选择器', '是', '禁用', '只读，不可修改'],
        ['成本分类', '单选下拉框', '是', '可编辑', '必选项'],
        ['成本池', '单选下拉框', '是', '可编辑', '必选项'],
        ['成本项名称', '文本输入框', '是', '可编辑', '必填，最大50字符'],
        ['地域', '单选下拉框', '是', '可编辑', '必选项'],
        ['公摊池', '单选下拉框', '否', '可编辑', '可选'],
        ['公摊类型', '单选下拉框', '是', '可编辑', '公摊/非公摊'],
        ['数据来源', '文本显示', '是', '禁用', '系统自动显示'],
        ['描述', '多行文本框', '否', '可编辑', '最大200字符'],
        ['单价', '数字输入框', '是', '可编辑', '必填，≥0，保留2位小数'],
        ['折旧月限', '数字输入框', '是', '可编辑', '必填，整数，1-600'],
        ['数量', '数字输入框', '是', '可编辑', '必填，≥0，整数'],
        ['单位', '文本输入框', '是', '可编辑', '必填，最大20字符'],
    ]
    add_table(doc, headers, rows)
    
    add_heading(doc, '10.5 资源详情Tab页', level=2)
    add_paragraph(doc, '资源详情Tab页用于配置成本项关联的资源规格信息。', bold=True)
    
    add_paragraph(doc, '10.5.1 资源规格表格', bold=True)
    headers = ['字段名称', '说明', '输入方式', '默认值']
    rows = [
        ['CPU', 'CPU核心数', '数字输入框', '0'],
        ['内存', '内存容量(GB)', '数字输入框', '0'],
        ['磁盘', '磁盘容量(GB)', '数字输入框', '0'],
    ]
    add_table(doc, headers, rows)
    
    add_paragraph(doc, '10.5.2 操作按钮', bold=True)
    headers = ['按钮', '功能', '位置']
    rows = [
        ['添加规格', '新增一行资源规格', '表格上方右侧'],
        ['删除', '删除当前行规格', '每行操作列'],
    ]
    add_table(doc, headers, rows)
    
    add_heading(doc, '10.6 计费资源Tab页', level=2)
    add_paragraph(doc, '计费资源Tab页用于配置成本项关联的计费资源信息。', bold=True)
    
    add_paragraph(doc, '10.6.1 计费资源表格', bold=True)
    headers = ['字段名称', '说明', '输入方式', '默认值']
    rows = [
        ['资源类型', '计费资源类型', '单选下拉框', '请选择'],
        ['资源名称', '资源实例名称', '文本输入框', '空'],
        ['资源数量', '资源实例数量', '数字输入框', '1'],
    ]
    add_table(doc, headers, rows)
    
    add_paragraph(doc, '10.6.2 操作按钮', bold=True)
    headers = ['按钮', '功能', '位置']
    rows = [
        ['添加资源', '新增一行计费资源', '表格上方右侧'],
        ['编辑', '编辑当前行资源', '每行操作列'],
        ['删除', '删除当前行资源', '每行操作列'],
    ]
    add_table(doc, headers, rows)
    
    add_heading(doc, '10.7 操作按钮区域', level=2)
    add_paragraph(doc, '页面底部固定操作栏：', bold=True)
    headers = ['按钮', '功能', '样式', '交互']
    rows = [
        ['保存', '保存所有修改', '主按钮（蓝色）', '校验通过后提交，成功后返回列表页'],
        ['取消', '放弃修改并返回', '次要按钮（灰色）', '弹出确认提示，确认后返回列表页'],
    ]
    add_table(doc, headers, rows)
    
    add_heading(doc, '10.8 编辑模式特性', level=2)
    add_paragraph(doc, '• 编辑按钮状态：进入编辑模式后，编辑按钮自动隐藏')
    add_paragraph(doc, '• 字段状态切换：点击编辑后，可编辑字段变为可输入状态')
    add_paragraph(doc, '• 实时校验：输入过程中实时校验字段合法性')
    add_paragraph(doc, '• 自动计算：单价、数量变化时自动计算折旧月价')
    add_paragraph(doc, '• 数据联动：资源规格变化时同步更新计费资源')
    
    add_heading(doc, '10.9 接口需求', level=2)
    add_paragraph(doc, '10.9.1 获取成本项详情', bold=True)
    add_paragraph(doc, '• 接口：GET /api/cost/detail/{id}')
    add_paragraph(doc, '• 返回：成本项完整信息（基本信息+资源详情+计费资源）')
    
    add_paragraph(doc, '10.9.2 保存成本项', bold=True)
    add_paragraph(doc, '• 接口：PUT /api/cost/{id}')
    add_paragraph(doc, '• 参数：完整成本项数据对象')
    add_paragraph(doc, '• 返回：操作结果')
    
    # 添加分页符
    doc.add_page_break()
    
    # ========== 历史月-详情页面 ==========
    add_heading(doc, '11. 历史月-详情页面需求', level=1)
    
    add_heading(doc, '11.1 页面概述', level=2)
    add_paragraph(doc, '历史月详情页面用于查看已归档的历史账期数据，仅支持查看，不支持任何编辑、添加、删除操作。')
    
    add_heading(doc, '11.2 页面入口', level=2)
    add_paragraph(doc, '• 入口路径：成本总览页面 → 历史月数据行 → 点击"详情"按钮')
    add_paragraph(doc, '• URL参数：cost-detail.html?id={id}&mode=view&month={month}')
    add_paragraph(doc, '• 标识特征：URL中包含mode=view参数')
    
    add_heading(doc, '11.3 页面布局', level=2)
    add_paragraph(doc, '页面采用左右分栏布局（同编辑页面）：')
    add_paragraph(doc, '• 左侧：成本项基本信息展示（固定宽度400px）')
    add_paragraph(doc, '• 右侧：Tab切换区域（资源详情/计费资源）')
    
    add_heading(doc, '11.4 历史月标识', level=2)
    add_paragraph(doc, '页面顶部显示历史月标识：', bold=True)
    add_paragraph(doc, '• 橙色徽章：页面标题旁显示"历史月"橙色徽章')
    add_paragraph(doc, '• 提示条：页面顶部显示警告提示条')
    add_paragraph(doc, '  - 内容："您正在查看历史成本数据，该数据已归档，仅支持查看"')
    add_paragraph(doc, '  - 样式：橙色背景，信息图标')
    
    add_heading(doc, '11.5 基本信息区域（只读模式）', level=2)
    add_paragraph(doc, '所有字段均为只读状态，以文本形式展示：', bold=True)
    
    headers = ['字段名称', '展示方式', '格式']
    rows = [
        ['成本编码', '纯文本', '灰色字体'],
        ['账期', '纯文本', '灰色字体'],
        ['成本分类', '纯文本', '黑色字体'],
        ['成本池', '纯文本', '黑色字体'],
        ['成本项名称', '纯文本', '黑色字体'],
        ['地域', '纯文本', '黑色字体'],
        ['公摊池', '纯文本', '黑色字体，无值显示"-"'],
        ['公摊类型', '纯文本', '黑色字体'],
        ['数据来源', '纯文本', '灰色字体'],
        ['描述', '纯文本', '黑色字体，多行文本自动换行'],
        ['单价', '纯文本', '¥XX,XXX.00'],
        ['折旧月限', '纯文本', 'XX个月'],
        ['数量', '纯文本', 'XX'],
        ['单位', '纯文本', '黑色字体'],
    ]
    add_table(doc, headers, rows)
    
    add_heading(doc, '11.6 资源详情Tab页（只读模式）', level=2)
    add_paragraph(doc, '资源详情以表格形式展示，所有字段不可编辑：', bold=True)
    
    headers = ['字段名称', '展示方式', '空值处理']
    rows = [
        ['CPU', '纯文本（XX核）', '显示"-"'],
        ['内存', '纯文本（XXGB）', '显示"-"'],
        ['磁盘', '纯文本（XXGB）', '显示"-"'],
    ]
    add_table(doc, headers, rows)
    
    add_paragraph(doc, '只读模式特性：', bold=True)
    add_paragraph(doc, '• 隐藏"添加规格"按钮')
    add_paragraph(doc, '• 隐藏每行的"删除"按钮')
    add_paragraph(doc, '• 表格行无hover效果')
    
    add_heading(doc, '11.7 计费资源Tab页（只读模式）', level=2)
    add_paragraph(doc, '计费资源以表格形式展示，所有字段不可编辑：', bold=True)
    
    headers = ['字段名称', '展示方式', '空值处理']
    rows = [
        ['资源类型', '纯文本', '显示"-"'],
        ['资源名称', '纯文本', '显示"-"'],
        ['资源数量', '纯文本', '显示"1"'],
    ]
    add_table(doc, headers, rows)
    
    add_paragraph(doc, '只读模式特性：', bold=True)
    add_paragraph(doc, '• 隐藏"添加资源"按钮')
    add_paragraph(doc, '• 隐藏每行的"编辑"按钮')
    add_paragraph(doc, '• 隐藏每行的"删除"按钮')
    add_paragraph(doc, '• 表格行无hover效果')
    
    add_heading(doc, '11.8 操作按钮区域', level=2)
    add_paragraph(doc, '历史月详情页面底部操作栏：', bold=True)
    headers = ['按钮', '功能', '样式']
    rows = [
        ['返回', '返回成本总览列表页', '次要按钮（灰色）'],
    ]
    add_table(doc, headers, rows)
    
    add_paragraph(doc, '注意：历史月详情页不显示"编辑"和"保存"按钮')
    
    add_heading(doc, '11.9 查看模式特性', level=2)
    add_paragraph(doc, '• 所有表单字段禁用：input、select、textarea均设置disabled属性')
    add_paragraph(doc, '• 所有按钮隐藏：添加、编辑、删除、保存按钮均不显示')
    add_paragraph(doc, '• 数据静态展示：以文本形式展示，无输入控件')
    add_paragraph(doc, '• 无操作反馈：无需校验、无需提示保存成功/失败')
    
    add_heading(doc, '11.10 与当前月编辑页对比', level=2)
    headers = ['对比项', '当前月-编辑页面', '历史月-详情页面']
    rows = [
        ['URL参数', 'mode=edit', 'mode=view'],
        ['账期字段', '禁用（灰色）', '只读（灰色文本）'],
        ['其他字段', '可编辑', '只读'],
        ['编辑按钮', '进入编辑后隐藏', '不显示'],
        ['保存按钮', '有', '无'],
        ['取消按钮', '有', '无（显示返回）'],
        ['添加资源按钮', '有', '无'],
        ['行级编辑按钮', '有', '无'],
        ['行级删除按钮', '有', '无'],
    ]
    add_table(doc, headers, rows)
    
    add_heading(doc, '11.11 接口需求', level=2)
    add_paragraph(doc, '11.11.1 获取成本项详情', bold=True)
    add_paragraph(doc, '• 接口：GET /api/cost/detail/{id}')
    add_paragraph(doc, '• 参数：month（历史月标识，用于数据归档校验）')
    add_paragraph(doc, '• 返回：成本项完整信息（只读）')
    
    add_paragraph(doc, '11.11.2 数据校验', bold=True)
    add_paragraph(doc, '• 后端校验：历史月数据不允许修改')
    add_paragraph(doc, '• 异常返回：尝试修改历史月数据时返回403错误')
    
    # 添加分页符
    doc.add_page_break()
    
    # ========== 页面状态流转 ==========
    add_heading(doc, '12. 页面状态流转说明', level=1)
    
    add_heading(doc, '12.1 状态定义', level=2)
    headers = ['状态', '标识', '说明']
    rows = [
        ['列表查看状态', 'cost-overview.html', '成本总览页面，查看成本列表'],
        ['当前月编辑状态', 'mode=edit', '编辑当前月成本项'],
        ['历史月查看状态', 'mode=view', '查看历史月成本项'],
    ]
    add_table(doc, headers, rows)
    
    add_heading(doc, '12.2 状态流转图', level=2)
    add_paragraph(doc, '成本总览列表页 → 点击当前月"编辑" → 当前月编辑页面（mode=edit）')
    add_paragraph(doc, '成本总览列表页 → 点击历史月"详情" → 历史月详情页面（mode=view）')
    add_paragraph(doc, '当前月编辑页面 → 点击"保存"/"取消" → 返回成本总览列表页')
    add_paragraph(doc, '历史月详情页面 → 点击"返回" → 返回成本总览列表页')
    
    add_heading(doc, '12.3 权限矩阵', level=2)
    headers = ['功能', '当前月编辑页', '历史月详情页']
    rows = [
        ['查看基本信息', '✓', '✓'],
        ['修改基本信息', '✓', '✗'],
        ['查看资源详情', '✓', '✓'],
        ['添加资源规格', '✓', '✗'],
        ['删除资源规格', '✓', '✗'],
        ['查看计费资源', '✓', '✓'],
        ['添加计费资源', '✓', '✗'],
        ['编辑计费资源', '✓', '✗'],
        ['删除计费资源', '✓', '✗'],
        ['保存修改', '✓', '✗'],
    ]
    add_table(doc, headers, rows)
    
    # 保存文档
    doc.save('成本管理-成本总览页面需求文档.docx')
    print("Word文档补充完成！")
    print("已添加以下内容：")
    print("1. 当前月-编辑页面需求（第10章）")
    print("2. 历史月-详情页面需求（第11章）")
    print("3. 页面状态流转说明（第12章）")

if __name__ == '__main__':
    main()
