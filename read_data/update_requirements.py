from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, size=10.5):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p

def add_table(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # 数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    return table

def update_document():
    # 打开现有文档
    doc = Document('成本管理-成本总览页面需求文档.docx')
    
    # 找到第10章和第11章的位置并删除旧内容，然后插入新内容
    # 由于python-docx不支持直接删除段落，我们采用追加新内容的方式
    
    # 添加分页符
    doc.add_page_break()
    
    # ========== 更新后的当前月-编辑页面需求 ==========
    add_heading(doc, '10. 当前月-编辑页面需求（已更新）', level=1)
    
    add_heading(doc, '10.1 页面概述', level=2)
    add_paragraph(doc, '当前月编辑页面用于编辑和修改当前账期的成本项信息，支持基本信息维护、资源详情查看和计费资源配置。页面根据数据来源区分可编辑字段范围。')
    
    add_heading(doc, '10.2 页面入口', level=2)
    add_paragraph(doc, '• 入口路径：成本总览页面 → 当前月数据行 → 点击"编辑"按钮')
    add_paragraph(doc, '• URL参数：cost-detail.html?id={id}&mode=edit')
    add_paragraph(doc, '• 页面标题：显示成本项名称（如"服务器A"）')
    
    add_heading(doc, '10.3 页面布局', level=2)
    add_paragraph(doc, '页面采用左右分栏布局：')
    add_paragraph(doc, '• 左侧：菜单导航栏（宽度220px，固定定位）')
    add_paragraph(doc, '• 右侧：主内容区（包含返回按钮、Tab切换区）')
    add_paragraph(doc, '• Tab页签：成本详情 | 资源详情 | 计费资源')
    
    add_heading(doc, '10.4 操作按钮区', level=2)
    add_paragraph(doc, '页面顶部左侧显示：', bold=True)
    headers = ['按钮', '功能', '样式']
    rows = [
        ['返回', '返回成本总览列表页', '次要按钮（灰色）'],
    ]
    add_table(doc, headers, rows)
    
    add_heading(doc, '10.5 成本详情Tab页', level=2)
    
    add_paragraph(doc, '10.5.1 基本信息模块', bold=True)
    add_paragraph(doc, '基本信息表单字段及编辑规则（根据数据来源区分）：')
    
    headers = ['字段名称', '控件类型', '是否必填', '编辑状态', '说明']
    rows = [
        ['成本分类', '单选下拉框', '是', '始终可编辑', '选项：设备费用、软件平台、网络带宽、运营费用、基础设施'],
        ['成本池', '单选下拉框', '是', '始终可编辑', '选项：通算服务器、智算服务器、存储、安全、网络、机房环境'],
        ['成本项', '文本输入框', '是', '始终可编辑', '成本项名称，placeholder="请输入成本项"'],
        ['地区', '单选下拉框', '否', '手动维护时可编辑', '选项：北京、上海、西安、总部'],
        ['公摊池', '单选下拉框', '是', '手动维护时可编辑', '选项：基础资源、容器、数据库、网络'],
        ['公摊类型', '单选下拉框', '是', '手动维护时可编辑', '选项：公摊成本、非公摊成本'],
        ['单价（元）', '数字输入框', '是', '手动维护时可编辑', 'step="0.01" min="0"，保留2位小数'],
        ['折旧月限', '数字输入框', '是', '手动维护时可编辑', 'step="1" min="1"，整数'],
        ['折旧月价', '数字输入框', '是', '手动维护时可编辑', '自动计算或手动输入'],
        ['数量', '数字输入框', '是', '手动维护时可编辑', 'step="1" min="1"，整数'],
        ['单位', '单选下拉框', '是', '手动维护时可编辑', '选项：台、套、Gbps、TB、月'],
        ['账期', '文本输入框', '是', '始终禁用', '格式：YYYY-MM，灰色背景，不可编辑'],
        ['月成本（元）', '数字输入框', '是', '手动维护时可编辑', 'step="0.01" min="0"'],
        ['公摊比例', '数字输入框', '是', '手动维护时可编辑', 'step="0.01" min="0" max="100"'],
        ['公摊后月总成本', '数字输入框', '是', '手动维护时可编辑', 'step="0.01" min="0"'],
        ['描述', '多行文本框', '否', '始终可编辑', 'rows="3"，placeholder="请输入描述"'],
    ]
    add_table(doc, headers, rows)
    
    add_paragraph(doc, '数据来源字段编辑规则：', bold=True)
    headers = ['数据来源', '可编辑字段', '不可编辑字段']
    rows = [
        ['资产系统-B类资产', '成本分类、成本池、成本项、公摊池、公摊类型、描述', '地区、单价、折旧月限、折旧月价、数量、单位、月成本、公摊比例、公摊后月总成本'],
        ['手动维护', '所有基本信息字段（除成本编码、数据来源、账期）', '成本编码、数据来源、账期'],
    ]
    add_table(doc, headers, rows)
    
    add_paragraph(doc, '10.5.2 系统信息模块', bold=True)
    add_paragraph(doc, '系统信息字段（全部只读，灰色背景）：')
    headers = ['字段名称', '控件类型', '说明']
    rows = [
        ['成本编码', '文本输入框', '系统自动生成，如COST-2026-0001'],
        ['数据来源', '单选下拉框', '选项：资产系统-B类资产、手动维护'],
        ['创建人', '文本输入框', '创建者用户名'],
        ['创建时间', '日期时间选择器', '格式：YYYY-MM-DDTHH:mm:ss'],
        ['更新人', '文本输入框', '最后更新者用户名'],
        ['更新时间', '日期时间选择器', '格式：YYYY-MM-DDTHH:mm:ss'],
    ]
    add_table(doc, headers, rows)
    
    add_paragraph(doc, '10.5.3 编辑/保存按钮逻辑', bold=True)
    add_paragraph(doc, '• 默认状态：显示"编辑"按钮（蓝色主按钮）')
    add_paragraph(doc, '• 点击编辑后：')
    add_paragraph(doc, '  - 隐藏"编辑"按钮')
    add_paragraph(doc, '  - 显示"保存"和"取消"按钮（位于Tab容器底部右侧）')
    add_paragraph(doc, '  - 根据数据来源启用相应字段')
    add_paragraph(doc, '  - 可编辑字段背景变为白色，cursor变为text')
    add_paragraph(doc, '• 点击保存后：')
    add_paragraph(doc, '  - 显示"保存成功"提示')
    add_paragraph(doc, '  - 显示"编辑"按钮')
    add_paragraph(doc, '  - 隐藏"保存/取消"按钮')
    add_paragraph(doc, '  - 禁用所有表单字段')
    add_paragraph(doc, '• 点击取消后：')
    add_paragraph(doc, '  - 显示"已取消修改"提示')
    add_paragraph(doc, '  - 恢复显示"编辑"按钮')
    add_paragraph(doc, '  - 隐藏"保存/取消"按钮')
    add_paragraph(doc, '  - 禁用所有表单字段')
    
    add_heading(doc, '10.6 资源详情Tab页', level=2)
    add_paragraph(doc, '资源详情以只读表格形式展示成本项关联的资源规格信息。')
    
    add_paragraph(doc, '10.6.1 资源规格表格', bold=True)
    headers = ['字段名称', '说明', '示例值']
    rows = [
        ['资源类型', '资源规格类型', 'CPU、内存、存储'],
        ['总量', '资源总量', '20（CPU核数）、40（内存GB）、500（存储GB）'],
        ['单位', '计量单位', '核、GB、GB'],
    ]
    add_table(doc, headers, rows)
    
    add_paragraph(doc, '10.6.2 表格特性', bold=True)
    add_paragraph(doc, '• 表格样式：白色背景，表头浅灰色背景（#fafafa）')
    add_paragraph(doc, '• 行hover效果：鼠标悬停时背景变为#f5f5f5')
    add_paragraph(doc, '• 无边框操作列：纯展示，无编辑/删除操作')
    
    add_heading(doc, '10.7 计费资源Tab页', level=2)
    add_paragraph(doc, '计费资源Tab页用于配置成本项关联的计费资源及单价计算公式。')
    
    add_paragraph(doc, '10.7.1 功能按钮', bold=True)
    headers = ['按钮', '功能', '位置']
    rows = [
        ['添加', '打开添加计费资源弹窗', '表格上方左侧'],
    ]
    add_table(doc, headers, rows)
    
    add_paragraph(doc, '10.7.2 计费资源表格', bold=True)
    headers = ['字段名称', '说明', '示例值']
    rows = [
        ['计费资源', '资源类型', 'CPU、内存、存储'],
        ['资源数量', '资源实例数量', '1'],
        ['单价计算公式', '计算公式及编辑入口', '基础费用 + 使用率 * 费率'],
        ['单价', '计算后的单价', '¥100'],
        ['单位', '计量单位', '核、GB、GB'],
    ]
    add_table(doc, headers, rows)
    
    add_paragraph(doc, '10.7.3 行级操作', bold=True)
    headers = ['按钮', '功能', '说明']
    rows = [
        ['编辑', '打开公式编辑器', '每行单价计算公式列内显示，小尺寸按钮（padding: 2px 8px, font-size: 12px）'],
    ]
    add_table(doc, headers, rows)
    
    add_paragraph(doc, '10.7.4 添加计费资源弹窗', bold=True)
    add_paragraph(doc, '• 触发方式：点击"添加"按钮')
    add_paragraph(doc, '• 弹窗标题：添加计费资源')
    add_paragraph(doc, '• 表单字段：')
    add_paragraph(doc, '  - 计费资源（必填）：下拉选择，选项：CPU、内存、存储')
    add_paragraph(doc, '• 操作按钮：取消（灰色）、确定（蓝色主按钮）')
    add_paragraph(doc, '• 校验规则：未选择资源时提示"请选择计费资源"')
    
    add_paragraph(doc, '10.7.5 单价计算公式编辑器', bold=True)
    add_paragraph(doc, '• 触发方式：点击计费资源行中的"编辑"按钮')
    add_paragraph(doc, '• 弹窗标题：编辑单价计算公式')
    add_paragraph(doc, '• 弹窗布局：')
    add_paragraph(doc, '  - 上半部分左右分栏：')
    add_paragraph(doc, '    - 左侧：变量列表（蓝色徽章按钮）')
    add_paragraph(doc, '      * 可选变量：单价、月成本、公摊后月总成本、CPU总量、内存总量、存储总量、流量总量')
    add_paragraph(doc, '    - 右侧：公式函数列表（蓝色徽章按钮）')
    add_paragraph(doc, '      * 可选函数：SUM()、AVERAGE()、MAX()、MIN()、IF()、ROUND()、ABS()')
    add_paragraph(doc, '  - 下半部分：')
    add_paragraph(doc, '    - 规则定义：多行文本输入框（高度120px，等宽字体）')
    add_paragraph(doc, '    - 辅助按钮：清空、测试')
    add_paragraph(doc, '    - 使用说明区域')
    add_paragraph(doc, '• 交互逻辑：点击变量/函数徽章自动添加到规则定义文本框')
    add_paragraph(doc, '• 操作按钮：取消（灰色）、保存（蓝色主按钮）')
    
    add_heading(doc, '10.8 新增模式（mode=add）', level=2)
    add_paragraph(doc, '当URL参数mode=add时，页面进入新增成本项模式：')
    add_paragraph(doc, '• 自动进入编辑状态：')
    add_paragraph(doc, '  - 隐藏"编辑"按钮')
    add_paragraph(doc, '  - 显示"保存/取消"按钮')
    add_paragraph(doc, '• 页面标题变为：新增成本项')
    add_paragraph(doc, '• 启用所有可编辑字段（账期始终禁用）')
    add_paragraph(doc, '• 表单默认值：')
    add_paragraph(doc, '  - 成本分类：设备费用')
    add_paragraph(doc, '  - 成本池：通算服务器')
    add_paragraph(doc, '  - 地区：北京')
    add_paragraph(doc, '  - 公摊池：基础资源')
    add_paragraph(doc, '  - 公摊类型：非公摊成本')
    add_paragraph(doc, '  - 账期：当前系统月份（格式：YYYY-MM）')
    add_paragraph(doc, '  - 折旧月限：36')
    add_paragraph(doc, '  - 数量：1')
    add_paragraph(doc, '  - 单位：台')
    add_paragraph(doc, '  - 其他字段：空值')
    
    add_heading(doc, '10.9 接口需求', level=2)
    add_paragraph(doc, '10.9.1 获取成本项详情', bold=True)
    add_paragraph(doc, '• 接口：GET /api/cost/detail/{id}')
    add_paragraph(doc, '• 返回：成本项完整信息（基本信息+资源详情+计费资源+系统信息）')
    
    add_paragraph(doc, '10.9.2 保存成本项', bold=True)
    add_paragraph(doc, '• 接口：PUT /api/cost/{id}')
    add_paragraph(doc, '• 参数：完整成本项数据对象')
    add_paragraph(doc, '• 返回：操作结果')
    
    add_paragraph(doc, '10.9.3 新增成本项', bold=True)
    add_paragraph(doc, '• 接口：POST /api/cost')
    add_paragraph(doc, '• 参数：成本项数据对象（无ID）')
    add_paragraph(doc, '• 返回：新增成本项ID')
    
    # 添加分页符
    doc.add_page_break()
    
    # ========== 更新后的历史月-详情页面需求 ==========
    add_heading(doc, '11. 历史月-详情页面需求（已更新）', level=2)
    
    add_heading(doc, '11.1 页面概述', level=2)
    add_paragraph(doc, '历史月详情页面用于查看已归档的历史账期数据，仅支持查看，不支持任何编辑、添加、删除操作。所有数据以只读形式展示。')
    
    add_heading(doc, '11.2 页面入口', level=2)
    add_paragraph(doc, '• 入口路径：成本总览页面 → 历史月数据行 → 点击"详情"按钮')
    add_paragraph(doc, '• URL参数：cost-detail.html?id={id}&mode=view&month={month}')
    add_paragraph(doc, '• 标识特征：URL中包含mode=view参数')
    add_paragraph(doc, '• 页面标题：成本项详情（历史月）')
    
    add_heading(doc, '11.3 页面布局', level=2)
    add_paragraph(doc, '页面布局与编辑页面相同，但所有内容均为只读：')
    add_paragraph(doc, '• 左侧：菜单导航栏（宽度220px）')
    add_paragraph(doc, '• 右侧：主内容区（包含返回按钮、Tab切换区）')
    add_paragraph(doc, '• Tab页签：成本详情 | 资源详情 | 计费资源')
    
    add_heading(doc, '11.4 历史月标识', level=2)
    add_paragraph(doc, '页面顶部显示历史月标识：')
    add_paragraph(doc, '• 页面标题显示：成本项详情（历史月）')
    add_paragraph(doc, '• 无橙色徽章和提示条（在列表页已显示）')
    
    add_heading(doc, '11.5 成本详情Tab页（只读模式）', level=2)
    
    add_paragraph(doc, '11.5.1 基本信息模块（全部只读）', bold=True)
    add_paragraph(doc, '所有字段均为禁用状态，灰色背景展示：')
    headers = ['字段名称', '展示方式', '样式']
    rows = [
        ['成本分类', '下拉框（禁用）', '灰色背景#f5f5f5'],
        ['成本池', '下拉框（禁用）', '灰色背景#f5f5f5'],
        ['成本项', '文本输入框（禁用）', '灰色背景#f5f5f5'],
        ['地区', '下拉框（禁用）', '灰色背景#f5f5f5'],
        ['公摊池', '下拉框（禁用）', '灰色背景#f5f5f5'],
        ['公摊类型', '下拉框（禁用）', '灰色背景#f5f5f5'],
        ['单价（元）', '数字输入框（禁用）', '灰色背景#f5f5f5'],
        ['折旧月限', '数字输入框（禁用）', '灰色背景#f5f5f5'],
        ['折旧月价', '数字输入框（禁用）', '灰色背景#f5f5f5'],
        ['数量', '数字输入框（禁用）', '灰色背景#f5f5f5'],
        ['单位', '下拉框（禁用）', '灰色背景#f5f5f5'],
        ['账期', '文本输入框（禁用）', '灰色背景#f5f5f5，cursor: not-allowed'],
        ['月成本（元）', '数字输入框（禁用）', '灰色背景#f5f5f5'],
        ['公摊比例', '数字输入框（禁用）', '灰色背景#f5f5f5'],
        ['公摊后月总成本', '数字输入框（禁用）', '灰色背景#f5f5f5'],
        ['描述', '多行文本框（禁用）', '灰色背景#f5f5f5'],
    ]
    add_table(doc, headers, rows)
    
    add_paragraph(doc, '11.5.2 系统信息模块（全部只读）', bold=True)
    add_paragraph(doc, '所有系统信息字段均为禁用状态：')
    headers = ['字段名称', '展示方式']
    rows = [
        ['成本编码', '文本输入框（禁用）'],
        ['数据来源', '下拉框（禁用）'],
        ['创建人', '文本输入框（禁用）'],
        ['创建时间', '日期时间选择器（禁用）'],
        ['更新人', '文本输入框（禁用）'],
        ['更新时间', '日期时间选择器（禁用）'],
    ]
    add_table(doc, headers, rows)
    
    add_paragraph(doc, '11.5.3 隐藏的操作按钮', bold=True)
    add_paragraph(doc, '以下按钮在历史月详情页不显示：')
    add_paragraph(doc, '• "编辑"按钮（完全隐藏）')
    add_paragraph(doc, '• "保存"按钮（完全隐藏）')
    add_paragraph(doc, '• "取消"按钮（完全隐藏）')
    
    add_heading(doc, '11.6 资源详情Tab页（只读模式）', level=2)
    add_paragraph(doc, '资源详情表格以只读形式展示：')
    headers = ['字段名称', '说明']
    rows = [
        ['资源类型', 'CPU、内存、存储'],
        ['总量', '资源数量'],
        ['单位', '计量单位（核、GB）'],
    ]
    add_table(doc, headers, rows)
    add_paragraph(doc, '特性：')
    add_paragraph(doc, '• 表格样式与编辑模式相同')
    add_paragraph(doc, '• 保留行hover效果')
    add_paragraph(doc, '• 无操作列')
    
    add_heading(doc, '11.7 计费资源Tab页（只读模式）', level=2)
    add_paragraph(doc, '计费资源表格以只读形式展示，隐藏所有操作按钮：')
    
    add_paragraph(doc, '11.7.1 隐藏的功能按钮', bold=True)
    add_paragraph(doc, '• "添加"按钮（id="add-billing-resource-btn"）完全隐藏')
    add_paragraph(doc, '• 通过JS设置：billingActionButtons.style.display = "none"')
    
    add_paragraph(doc, '11.7.2 计费资源表格（只读）', bold=True)
    headers = ['字段名称', '说明']
    rows = [
        ['计费资源', 'CPU、内存、存储'],
        ['资源数量', '固定为1'],
        ['单价计算公式', '显示公式文本，无编辑按钮'],
        ['单价', '计算后的单价'],
        ['单位', '计量单位'],
    ]
    add_table(doc, headers, rows)
    
    add_paragraph(doc, '11.7.3 隐藏的行级操作按钮', bold=True)
    add_paragraph(doc, '• "编辑"按钮（class="billing-edit-btn"）完全隐藏')
    add_paragraph(doc, '• 通过JS设置：billingEditBtns.forEach(btn => btn.style.display = "none")')
    add_paragraph(doc, '• 单价计算公式列仅显示文本，无编辑入口')
    
    add_heading(doc, '11.8 操作按钮区', level=2)
    add_paragraph(doc, '历史月详情页面仅显示：')
    headers = ['按钮', '功能', '位置']
    rows = [
        ['返回', '返回成本总览列表页', '页面顶部左侧'],
    ]
    add_table(doc, headers, rows)
    
    add_heading(doc, '11.9 查看模式实现逻辑（mode=view）', level=2)
    add_paragraph(doc, '当URL参数mode=view时，页面通过JavaScript执行以下操作：')
    add_paragraph(doc, '1. 隐藏编辑按钮：')
    add_paragraph(doc, '   document.getElementById("edit-btn").style.display = "none"')
    add_paragraph(doc, '2. 隐藏保存/取消按钮区：')
    add_paragraph(doc, '   document.getElementById("save-cancel-buttons").style.display = "none"')
    add_paragraph(doc, '3. 隐藏计费资源添加按钮：')
    add_paragraph(doc, '   document.getElementById("billing-action-buttons").style.display = "none"')
    add_paragraph(doc, '4. 隐藏所有计费资源编辑按钮：')
    add_paragraph(doc, '   document.querySelectorAll(".billing-edit-btn").forEach(btn => btn.style.display = "none")')
    add_paragraph(doc, '5. 更新页面标题：')
    add_paragraph(doc, '   document.querySelector(".page-title").textContent = "成本项详情（历史月）"')
    
    add_heading(doc, '11.10 与当前月编辑页对比', level=2)
    headers = ['对比项', '当前月-编辑页面', '历史月-详情页面']
    rows = [
        ['URL参数', 'mode=edit 或 mode=add', 'mode=view'],
        ['页面标题', '成本项名称 / 新增成本项', '成本项详情（历史月）'],
        ['编辑按钮', '有（默认显示，点击后隐藏）', '无（完全隐藏）'],
        ['保存/取消按钮', '编辑模式下显示', '完全隐藏'],
        ['添加计费资源按钮', '有', '无（完全隐藏）'],
        ['计费资源编辑按钮', '每行显示', '完全隐藏'],
        ['基本信息字段', '部分可编辑（根据数据来源）', '全部禁用'],
        ['系统信息字段', '全部禁用', '全部禁用'],
        ['资源详情表格', '只读展示', '只读展示'],
        ['计费资源表格', '可添加、可编辑公式', '纯展示'],
        ['公式编辑器弹窗', '可打开编辑', '不可打开'],
        ['添加资源弹窗', '可打开添加', '不可打开'],
    ]
    add_table(doc, headers, rows)
    
    add_heading(doc, '11.11 接口需求', level=2)
    add_paragraph(doc, '11.11.1 获取成本项详情', bold=True)
    add_paragraph(doc, '• 接口：GET /api/cost/detail/{id}')
    add_paragraph(doc, '• 参数：month（历史月标识）')
    add_paragraph(doc, '• 返回：成本项完整信息（只读）')
    
    add_paragraph(doc, '11.11.2 数据安全校验', bold=True)
    add_paragraph(doc, '• 后端校验：历史月数据不允许修改')
    add_paragraph(doc, '• 异常返回：尝试修改历史月数据时返回403 Forbidden错误')
    add_paragraph(doc, '• 错误信息："历史数据不可编辑"')
    
    add_heading(doc, '11.12 页面加载完成后的特殊处理', level=2)
    add_paragraph(doc, '页面DOM加载完成后，确保账期字段始终不可编辑：')
    add_paragraph(doc, '```javascript')
    add_paragraph(doc, 'document.addEventListener("DOMContentLoaded", function() {')
    add_paragraph(doc, '    const costMonthField = document.getElementById("cost-month");')
    add_paragraph(doc, '    if (costMonthField) {')
    add_paragraph(doc, '        costMonthField.disabled = true;')
    add_paragraph(doc, '        costMonthField.style.backgroundColor = "#f5f5f5";')
    add_paragraph(doc, '        costMonthField.style.cursor = "not-allowed";')
    add_paragraph(doc, '    }')
    add_paragraph(doc, '});')
    add_paragraph(doc, '```')
    
    # 保存文档（使用新文件名避免占用冲突）
    output_path = '成本管理-成本总览页面需求文档_v2.0.docx'
    doc.save(output_path)
    print(f"Word文档更新完成！")
    print(f"保存路径: {output_path}")
    print("\n已根据cost-detail.html实际代码更新以下内容：")
    print("\n第10章 当前月-编辑页面需求：")
    print("- 补充了数据来源区分可编辑字段的详细规则")
    print("- 补充了系统信息模块（成本编码、创建人等）")
    print("- 补充了编辑/保存/取消按钮的详细交互逻辑")
    print("- 补充了资源详情Tab页的表格结构")
    print("- 补充了计费资源Tab页的完整功能（添加按钮、表格、公式编辑器）")
    print("- 补充了新增模式（mode=add）的详细说明")
    print("- 补充了公式编辑器弹窗的详细设计")
    print("\n第11章 历史月-详情页面需求：")
    print("- 补充了所有字段的只读展示方式")
    print("- 补充了隐藏按钮的详细列表（编辑、保存、添加、公式编辑）")
    print("- 补充了查看模式的JavaScript实现逻辑")
    print("- 补充了与当前月编辑页的详细对比表格")
    print("- 补充了账期字段的特殊处理逻辑")

if __name__ == '__main__':
    update_document()
